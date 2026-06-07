from __future__ import annotations

from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
OUT_DIR = ROOT / "AgentShield_Production_Documentation_Pack"
DOCX_DIR = OUT_DIR / "docx"
MD_DIR = OUT_DIR / "sources"


@dataclass
class Block:
    kind: str
    text: str | None = None
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class DocSpec:
    filename: str
    title: str
    subtitle: str
    blocks: list[Block]


def h(level: int, text: str) -> Block:
    return Block(kind=f"h{level}", text=text)


def p(text: str) -> Block:
    return Block(kind="p", text=text)


def bullets(items: Iterable[str]) -> Block:
    return Block(kind="bullets", rows=[[item] for item in items])


def numbered(items: Iterable[str]) -> Block:
    return Block(kind="numbered", rows=[[item] for item in items])


def table(rows: list[list[str]]) -> Block:
    return Block(kind="table", rows=rows)


def callout(text: str) -> Block:
    return Block(kind="callout", text=text)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 18, "0F172A"),
        ("Heading 2", 13.5, "1E3A8A"),
        ("Heading 3", 11.5, "374151"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0F172A")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub_run = sub.add_run(subtitle)
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.color.rgb = RGBColor.from_string("475569")

    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Production documentation pack | Version 1.0 | {date.today().isoformat()}")
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor.from_string("64748B")


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    t = doc.add_table(rows=len(rows), cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = t.cell(r_idx, c_idx)
            cell.text = row[c_idx] if c_idx < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8.6)
                    if r_idx == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string("F8FAFC")
            if r_idx == 0:
                set_cell_shading(cell, "1E3A8A")
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    doc.add_paragraph()


def build_docx(spec: DocSpec) -> None:
    doc = Document()
    style_document(doc)
    add_title(doc, spec.title, spec.subtitle)
    for block in spec.blocks:
        if block.kind == "h1":
            doc.add_heading(block.text or "", level=1)
        elif block.kind == "h2":
            doc.add_heading(block.text or "", level=2)
        elif block.kind == "h3":
            doc.add_heading(block.text or "", level=3)
        elif block.kind == "p":
            doc.add_paragraph(block.text or "")
        elif block.kind == "callout":
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.15)
            run = para.add_run(block.text or "")
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("1E3A8A")
        elif block.kind == "bullets":
            for row in block.rows:
                doc.add_paragraph(row[0], style="List Bullet")
        elif block.kind == "numbered":
            for row in block.rows:
                doc.add_paragraph(row[0], style="List Number")
        elif block.kind == "table":
            add_table(doc, block.rows)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{spec.title} | AgentShield Production Documentation Pack"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("64748B")
    doc.save(DOCX_DIR / spec.filename)


def block_to_md(block: Block) -> str:
    if block.kind.startswith("h"):
        level = int(block.kind[1])
        return f"{'#' * level} {block.text}\n"
    if block.kind in {"p", "callout"}:
        return f"{block.text}\n"
    if block.kind == "bullets":
        return "\n".join(f"- {row[0]}" for row in block.rows) + "\n"
    if block.kind == "numbered":
        return "\n".join(f"{i + 1}. {row[0]}" for i, row in enumerate(block.rows)) + "\n"
    if block.kind == "table":
        rows = block.rows
        if not rows:
            return ""
        cols = max(len(row) for row in rows)
        normalized = [row + [""] * (cols - len(row)) for row in rows]
        header = "| " + " | ".join(normalized[0]) + " |"
        sep = "| " + " | ".join("---" for _ in range(cols)) + " |"
        body = ["| " + " | ".join(row) + " |" for row in normalized[1:]]
        return "\n".join([header, sep] + body) + "\n"
    return ""


def build_md(spec: DocSpec) -> None:
    path = MD_DIR / spec.filename.replace(".docx", ".md")
    lines = [f"# {spec.title}", "", spec.subtitle, "", f"Version 1.0 | {date.today().isoformat()}", ""]
    for block in spec.blocks:
        lines.append(block_to_md(block))
    path.write_text("\n".join(lines), encoding="utf-8")


MASTER = DocSpec(
    filename="AgentShield_Master_Production_Plan.docx",
    title="AgentShield Master Production Plan",
    subtitle="Unified product, architecture, UX, security, implementation, rollout, and acceptance plan.",
    blocks=[
        h(1, "1. Executive Summary"),
        p("AgentShield is a production-grade security middleware and observability platform for autonomous and multi-agent LLM systems. It protects the runtime boundary where agents receive messages, exchange handoffs, and invoke tools. The system combines deterministic policy enforcement, cryptographic agent identity, prompt-injection detection, append-only audit logging, trust scoring, and a live operator dashboard."),
        p("This plan supersedes the prior PRD, TRD, and design document as the authoritative build source. The original documents remain historical inputs. Their strongest ideas are carried forward: real-time threat detection, identity verification, permission enforcement, hash-chained audit proof, 3D trust visualization, and an attack simulation demo. Their contradictions are resolved here into one production v1 plan."),
        callout("Production v1 principle: block high-confidence unsafe behavior synchronously, log every protected decision, and run slow AI-assisted enrichment asynchronously unless the caller explicitly requests deep analysis."),
        h(1, "2. Goals And Success Criteria"),
        table([
            ["Goal", "Production Success Criteria", "Verification"],
            ["Runtime security control", "Protected calls are evaluated before execution and unauthorized calls are blocked by default.", "Integration tests for message, tool-call, and inter-agent paths."],
            ["Provable auditability", "Every protected decision writes an immutable ledger entry with a verifiable hash chain.", "Tamper test breaks verification at the expected entry."],
            ["Practical latency", "P95 synchronous protection path under 200 ms without LLM calls; deep AI enrichment is async or explicitly requested.", "Load test on analyze and tool-call endpoints."],
            ["Operator visibility", "Dashboard shows live events, active agents, trust score changes, blocked events, and ledger status.", "Browser smoke test and WebSocket event replay."],
            ["Agent integration", "Python SDK can spawn an agent, analyze messages, check tool calls, and fetch verdicts using documented contracts.", "SDK example runs end to end against local and deployed backend."],
        ]),
        h(1, "3. Users And Use Cases"),
        table([
            ["Persona", "Need", "AgentShield Outcome"],
            ["AI security engineer", "Catch prompt injection, spoofing, and unsafe tool use before production damage.", "Drop-in middleware and dashboard with policy, audit, and detection evidence."],
            ["AI infrastructure engineer", "Add protection without rebuilding existing LangChain/OpenAI agent flows.", "SDK and REST contracts that wrap existing message and tool-call paths."],
            ["Security reviewer", "Verify what an autonomous agent did and why a decision was allowed or blocked.", "Immutable ledger, evidence fields, verdict details, and verification endpoint."],
            ["Product/demo evaluator", "See a real working system with visible security behavior.", "Attack simulator, live event feed, and 3D network visualization backed by real APIs."],
        ]),
        h(1, "4. Production Scope"),
        table([
            ["Priority", "Included Capabilities", "Deferred Capabilities"],
            ["P0 Core", "API key auth, RS256 agent JWT, deterministic injection checks, permission enforcement, trust score updates, ledger writes, ledger verification, Python SDK basics.", "Billing, SSO, tenant admin UI, long-term threat intelligence marketplace."],
            ["P1 Operator Product", "Dashboard, ledger browser, attack simulator, WebSocket stream, agent registry, deployment runbook, demo agent.", "Full visual no-code policy builder, advanced analytics, mobile-native app."],
            ["P2 Advanced Intelligence", "Async LLM enrichment, ReAct trace for demos, vector-backed threat context, adaptive risk recommendations.", "Autonomous remediation without human approval, production sandbox execution isolation."],
        ]),
        p("The production v1 build must not depend on an LLM call to block obvious attacks, invalid identity, or unauthorized tools. LLM and ReAct behavior is useful for explanation, enrichment, and demo traces, but not for the critical enforcement path."),
        h(1, "5. Resolved Architecture"),
        table([
            ["Layer", "Decision", "Reason"],
            ["Client and SDK", "Python SDK and direct REST clients authenticate with X-AgentShield-API-Key.", "Separates customer/client authorization from agent identity."],
            ["Agent identity", "Each spawned agent receives a signed RS256 JWT with agent_id, tenant_id, scopes, iat, exp, jti, and key id.", "Prevents spoofed inter-agent and tool-call requests."],
            ["Synchronous guard", "Fast policy engine verifies API key, JWT, permission manifest, pattern classifier, trust state, and ledger write.", "Meets latency target and blocks deterministic risks."],
            ["Async enrichment", "Queue worker runs optional LLM classification, threat enrichment, vector search, and narrative explanation.", "Preserves UX without slowing protected execution."],
            ["Data", "PostgreSQL stores agents, API keys, permission manifests, audit ledger, threats, trust history, revoked tokens, and event outbox.", "Keeps the production v1 stack deployable and inspectable."],
            ["Frontend", "Next.js dashboard consumes REST and WebSocket APIs; 3D network reads stable agent/event payloads.", "Prevents frontend from relying on mocked visual-only state."],
        ]),
        h(1, "6. Core Runtime Flow"),
        numbered([
            "Client sends a protected request with X-AgentShield-API-Key and, when acting as an agent, an RS256 agent JWT.",
            "Auth middleware validates API key status, tenant scope, rate limit bucket, and request body size.",
            "Identity verifier validates JWT signature, exp, issuer, audience, jti revocation, tenant match, and agent status.",
            "Policy engine evaluates message or tool-call permissions using deny-by-default manifests.",
            "Injection detector runs deterministic pattern checks and optionally flags the event for async LLM enrichment.",
            "Trust engine computes the delta and stores current score plus a trust history row in the same transaction.",
            "Ledger writer appends a hash-chained audit row using a locked previous-head read and immutable DB role.",
            "Event outbox emits a WebSocket event after commit; dashboard updates live counters and feeds.",
        ]),
        h(1, "7. Frontend And UX Plan"),
        p("The production dashboard is the first product surface. The landing page is secondary and may reuse dashboard metrics, but the system should open directly into operational value: active agents, blocked threats, trust changes, and ledger health."),
        table([
            ["Route", "Purpose", "Must Show Real Data"],
            ["/dashboard", "Primary console with 3D network, live feed, metrics strip, and agent side panel.", "Agents, events, trust scores, WebSocket status."],
            ["/ledger", "Audit ledger browser with verify-chain action and tamper-demo explanation.", "Ledger rows, hashes, verification result."],
            ["/attack-sim", "Security testing terminal with curated attacks and custom payloads.", "Attack result, verdict, evidence, latency, ledger id."],
            ["/agents", "Agent registry, status, permissions, trust history, revoke action.", "Agent records, manifests, trust history."],
            ["/", "Production landing page with live counters and path into console.", "At minimum threat count and system status."],
        ]),
        p("Mobile behavior must preserve utility. Under 768 px, replace the full Three.js graph with a sortable agent/event list and compact trust indicators. Do not use a static screenshot as the only mobile fallback."),
        h(1, "8. Implementation Phases"),
        table([
            ["Phase", "Deliverable", "Exit Criteria"],
            ["0 - Foundation", "Repo scaffold, env template, database migrations, API health, CI checks.", "Local backend and frontend boot; migrations apply cleanly."],
            ["1 - Security Core", "API key auth, agent JWTs, permission manifests, deterministic injection detector.", "Protected message and tool-call tests pass."],
            ["2 - Ledger And Trust", "Immutable audit ledger, chain verification, trust scoring, event outbox.", "Tamper test and concurrent ledger-write test pass."],
            ["3 - SDK And Demo Agent", "Python SDK, demo protected agent, attack fixtures.", "SDK example blocks known attacks and logs ledger entries."],
            ["4 - Dashboard", "Next.js dashboard, ledger browser, attack simulator, WebSocket feed.", "Browser smoke test covers core routes using real backend."],
            ["5 - Production Readiness", "Deployment, monitoring, rate limits, backups, key rotation, incident runbooks.", "Production smoke test and readiness checklist pass."],
        ]),
        h(1, "9. Requirement Disposition"),
        table([
            ["Original Requirement", "Disposition", "Production v1 Decision"],
            ["Prompt injection detector", "Carried forward and revised.", "Deterministic fast path required; LLM classifier async or explicit deep mode."],
            ["Agent identity tokens", "Carried forward.", "RS256 JWT with jti, kid, tenant_id, scopes, expiry, revocation, and rotation."],
            ["Permission enforcement", "Carried forward and tightened.", "Deny by default; server-side enforcement before tool execution."],
            ["Tamper-proof audit ledger", "Carried forward and strengthened.", "DB-level append-only controls, hash chain, transaction locking, verify API."],
            ["Threat dashboard and 3D graph", "Carried forward as P1.", "Uses stable REST/WebSocket data; mobile gets functional list fallback."],
            ["Attack simulation sandbox", "Carried forward as P1.", "Runs curated payloads against the real analyze endpoint and writes ledger entries."],
            ["Trust score engine", "Carried forward and specified.", "Bounded score from 0.0 to 1.0 with deterministic event deltas."],
            ["REST API and SDK", "Carried forward.", "SDK wraps documented endpoints and never hides verdict errors."],
            ["Alert system", "Carried forward as P1/P2.", "Severity events now; resolution workflow later."],
            ["Multiple specialist agents and vector store", "Deferred to P2.", "Not required for production core; can enrich threat context later."],
        ]),
        h(1, "10. Risks And Controls"),
        table([
            ["Risk", "Control"],
            ["LLM latency breaks protection path.", "Keep synchronous protection deterministic; enqueue enrichment."],
            ["Ledger race conditions break hash chain.", "Use serializable transaction or advisory lock for append; test concurrent writes."],
            ["API key or JWT leakage.", "Hash API keys at rest, short-lived agent tokens, revocation, rotation, scoped permissions."],
            ["Dashboard becomes visual-only demo.", "All dashboard components must bind to real API contracts or documented fixtures from the backend."],
            ["Overbroad agent permissions.", "Deny by default; require explicit manifest actions; surface risky manifests in UI."],
        ]),
    ],
)


CONTRACTS = DocSpec(
    filename="AgentShield_API_and_Event_Contracts.docx",
    title="AgentShield API And Event Contracts",
    subtitle="Authoritative schemas for REST, WebSocket, SDK, database, errors, and event payloads.",
    blocks=[
        h(1, "1. Contract Rules"),
        bullets([
            "All production endpoints require X-AgentShield-API-Key unless explicitly marked public.",
            "Agent-acting endpoints also require an RS256 agent JWT in Authorization: Bearer <token>.",
            "Every protected decision returns a verdict and writes a ledger entry before returning success.",
            "All timestamps are ISO 8601 UTC strings. All IDs are UUID strings unless noted.",
            "All errors use the standard error envelope defined in this document.",
        ]),
        h(1, "2. Authentication Headers"),
        table([
            ["Header", "Required For", "Format", "Validation"],
            ["X-AgentShield-API-Key", "All non-public REST and WebSocket handshakes", "as_live_<token>", "Hash lookup, tenant active, scope allowed, rate limit not exceeded."],
            ["Authorization", "Agent message, tool-call, trust, and agent-specific requests", "Bearer <RS256 JWT>", "Signature, kid, issuer, audience, exp, jti, tenant_id, agent_id, status."],
            ["X-Request-Id", "Recommended all requests", "Client UUID", "Echoed in response and event metadata."],
        ]),
        h(1, "3. Core REST Endpoints"),
        table([
            ["Method", "Path", "Purpose", "Auth", "Success Response"],
            ["POST", "/v1/agents", "Spawn/register protected agent.", "API key", "AgentResponse"],
            ["GET", "/v1/agents", "List agents for tenant.", "API key", "AgentListResponse"],
            ["POST", "/v1/agents/{agent_id}/revoke", "Revoke agent and active JWT ids.", "API key", "AgentResponse"],
            ["POST", "/v1/shield/analyze", "Evaluate an incoming or outgoing message.", "API key + JWT", "SecurityVerdict"],
            ["POST", "/v1/shield/tool-call", "Authorize a tool call before execution.", "API key + JWT", "SecurityVerdict"],
            ["GET", "/v1/ledger", "Paginated audit ledger.", "API key", "LedgerPage"],
            ["GET", "/v1/ledger/verify", "Verify ledger hash chain.", "API key", "LedgerVerification"],
            ["GET", "/v1/threats", "List threat events.", "API key", "ThreatPage"],
            ["POST", "/v1/attack-sim/run", "Run curated or custom attack against sandbox agent.", "API key", "AttackSimulationResult"],
            ["GET", "/health", "Service health for deploy checks.", "Public", "HealthResponse"],
        ]),
        h(1, "4. Request And Response Schemas"),
        table([
            ["Schema", "Fields"],
            ["AgentCreateRequest", "name: string; type: user_agent|research_agent|executor_agent|security_agent|custom; permissions: PermissionManifest; metadata?: object"],
            ["AgentResponse", "agent_id: uuid; tenant_id: uuid; name: string; status: active|suspended|revoked; trust_score: number; token: string; token_expires_at: datetime; permissions: PermissionManifest"],
            ["PermissionManifest", "tools: object where each key is a tool name and value is allowed action array; default_action: deny; max_risk_level?: low|medium|high"],
            ["AnalyzeRequest", "agent_id: uuid; direction: inbound|outbound|inter_agent; message: string; context?: object; deep_analysis?: boolean"],
            ["ToolCallRequest", "agent_id: uuid; tool_name: string; action: string; arguments_hash?: string; risk_context?: object"],
            ["SecurityVerdict", "allowed: boolean; verdict: ALLOWED|BLOCKED|FLAGGED; threat_level: NONE|LOW|MEDIUM|HIGH|CRITICAL; reason: string; evidence: Evidence[]; trust_delta: number; trust_score_after: number; ledger_id: integer; latency_ms: integer; async_enrichment_id?: uuid"],
            ["Evidence", "source: pattern|permission|identity|trust|llm|manual; code: string; message: string; confidence?: number; span?: string"],
            ["LedgerEntry", "id: integer; tenant_id: uuid; agent_id?: uuid; event_type: message|tool_call|inter_agent|auth|system; severity: INFO|WARN|CRITICAL; verdict: ALLOWED|BLOCKED|FLAGGED; event_data: object; prev_hash: string; curr_hash: string; created_at: datetime"],
            ["ErrorEnvelope", "error: {code: string; message: string; request_id?: string; details?: object}"],
        ]),
        h(1, "5. Error Codes"),
        table([
            ["Code", "HTTP", "Meaning"],
            ["AUTH_API_KEY_MISSING", "401", "X-AgentShield-API-Key was not provided."],
            ["AUTH_API_KEY_INVALID", "401", "API key hash not found or inactive."],
            ["AUTH_AGENT_TOKEN_INVALID", "401", "JWT signature, issuer, audience, expiry, or jti failed validation."],
            ["AUTH_AGENT_TOKEN_REVOKED", "401", "JWT jti or agent status is revoked."],
            ["POLICY_TOOL_DENIED", "403", "Requested tool or action is not in the permission manifest."],
            ["POLICY_MESSAGE_BLOCKED", "403", "Message was blocked by injection or trust policy."],
            ["LEDGER_APPEND_FAILED", "500", "Ledger append transaction failed; protected decision must not proceed silently."],
            ["RATE_LIMITED", "429", "Tenant or API key exceeded configured limit."],
        ]),
        h(1, "6. WebSocket Events"),
        table([
            ["Event", "Payload Fields", "Use"],
            ["security.event.created", "event_id, ledger_id, agent_id, event_type, severity, verdict, threat_level, created_at", "Live feed and counters."],
            ["agent.trust.updated", "agent_id, previous_score, delta, score_after, reason, created_at", "Node color and trust chart."],
            ["ledger.verification.completed", "valid, entries_checked, broken_at, checked_at", "Ledger page verify UI."],
            ["attack_sim.completed", "simulation_id, attack_type, detected, verdict, latency_ms, ledger_id", "Attack simulation results."],
            ["system.health.changed", "component, status, message, created_at", "Status bar and operational alerts."],
        ]),
        h(1, "7. Database Tables"),
        table([
            ["Table", "Purpose", "Important Constraints"],
            ["tenants", "Customer/workspace boundary.", "id primary key; status active/suspended."],
            ["api_keys", "Client credentials.", "token_hash only; never store raw key; scopes; last_used_at."],
            ["agents", "Agent registry and current trust.", "tenant_id foreign key; status enum; trust_score 0.0 to 1.0."],
            ["agent_tokens", "JWT jti registry and revocation.", "jti unique; expires_at; revoked_at nullable."],
            ["permission_manifests", "Versioned permissions.", "agent_id; version; manifest jsonb; active flag."],
            ["audit_ledger", "Immutable proof log.", "curr_hash unique; no update/delete role; trigger rejects mutation."],
            ["threat_events", "Queryable incidents.", "ledger_id foreign key; attack_type; confidence; resolved state."],
            ["trust_history", "Trust score changes.", "score_after check 0.0 to 1.0; ledger_id link."],
            ["event_outbox", "Reliable WebSocket/event dispatch.", "processed_at nullable; retry_count."],
        ]),
        h(1, "8. SDK Contract"),
        table([
            ["Method", "Input", "Output"],
            ["spawn_agent", "name, permissions, metadata", "Agent object with id, token, expires_at."],
            ["analyze", "agent_id, token, message, direction, context, deep_analysis", "SecurityVerdict; raises SecurityBlocked when allowed is false if configured."],
            ["check_tool_call", "agent_id, token, tool_name, action, arguments", "SecurityVerdict."],
            ["verify_ledger", "optional start/end ids", "LedgerVerification."],
            ["run_attack_sim", "attack_type or payload", "AttackSimulationResult."],
        ]),
        h(1, "9. Contract Invariants"),
        bullets([
            "A BLOCKED verdict is final for the protected request. Async enrichment may add explanation, but it must not convert a blocked action into an allowed action without a separate audited admin override.",
            "A FLAGGED verdict means execution may continue only when the caller has explicitly configured allow_on_flagged for that route. The default is to stop and require operator review.",
            "Every SecurityVerdict must include ledger_id. If the ledger write fails, the protected request must fail closed with LEDGER_APPEND_FAILED.",
            "Frontend code must treat unknown enum values as unsafe: unknown severity renders as WARN, unknown verdict renders as BLOCKED, and unknown event_type renders as system.",
            "SDK clients must expose raw response bodies for debugging but must redact API keys and JWTs in exception messages.",
        ]),
        h(1, "10. Example Protected Tool-Call Request"),
        table([
            ["Field", "Example", "Notes"],
            ["agent_id", "6d4f8a7e-2f19-4a4e-8bd2-9d9c7e4f8a10", "Must match JWT agent_id claim."],
            ["tool_name", "web_search", "Exact registered tool name, not display label."],
            ["action", "read", "Action checked against permission manifest."],
            ["arguments_hash", "sha256:ab12...", "Hash avoids logging sensitive arguments while preserving auditability."],
            ["risk_context", "{\"destination\":\"external\"}", "Optional structured context for policy and audit."],
        ]),
    ],
)


SECURITY = DocSpec(
    filename="AgentShield_Security_Threat_Model.docx",
    title="AgentShield Security Threat Model",
    subtitle="Trust boundaries, abuse cases, mitigations, ledger guarantees, auth model, and residual risks.",
    blocks=[
        h(1, "1. Security Objective"),
        p("AgentShield protects autonomous agent systems from malicious instructions, identity spoofing, unauthorized tool use, and unverifiable behavior. It is a guardrail and audit layer, not a sandbox or full execution isolation system."),
        h(1, "2. Trust Boundaries"),
        table([
            ["Boundary", "Trusted Side", "Untrusted Side", "Controls"],
            ["Client to AgentShield API", "AgentShield backend", "SDK callers, browsers, external services", "API key auth, TLS, rate limits, body limits, request logging."],
            ["Agent identity", "JWT issuer and verifier", "Agent runtime, copied tokens, forged IDs", "RS256 signatures, jti registry, short expiry, revocation."],
            ["Tool execution", "Permission engine", "Agent-requested tool name/action/arguments", "Deny-by-default manifest, server-side check before execution."],
            ["Ledger", "Append transaction and verifier", "Application bugs, operators, compromised clients", "DB role separation, mutation trigger, hash chain, backups."],
            ["LLM enrichment", "Deterministic verdict already made", "Model output and external retrieved context", "Async processing, schema validation, no authority to override blocking decision automatically."],
        ]),
        h(1, "3. Primary Threats"),
        table([
            ["Threat", "Attack Example", "Mitigation", "Residual Risk"],
            ["Prompt injection", "User asks agent to ignore prior instructions or reveal secrets.", "Pattern classifier, semantic enrichment, block high-confidence signatures, evidence capture.", "Novel attacks can evade patterns; dataset must evolve."],
            ["Identity spoofing", "Request claims to be a trusted agent using another agent_id.", "JWT subject/agent_id match, tenant match, signature verification, revocation.", "Stolen token works until expiry unless detected and revoked."],
            ["Unauthorized tool use", "Research agent requests db_write or external transfer.", "Deny-by-default permission manifests and action-level checks.", "Badly configured manifests can still allow excessive scope."],
            ["Ledger tampering", "Operator updates event_data after a bad decision.", "No update/delete role, mutation trigger, hash verification, offsite backups.", "Superuser database compromise can bypass app controls."],
            ["Replay attack", "Old allowed verdict is reused for a new action.", "JWT expiry, jti registry, request ids, optional nonce for high-risk tools.", "Replay prevention depends on client adoption for nonces."],
            ["Dashboard deception", "Frontend displays mocked security state.", "Dashboard must consume real APIs or named fixtures in demo mode.", "Demo mode must be visibly labeled."],
        ]),
        h(1, "4. Auth And Key Rotation"),
        bullets([
            "API keys are generated once, shown once, and stored as salted hashes with prefix metadata.",
            "Each API key has tenant_id, scopes, status, created_at, last_used_at, expires_at, and optional allowed origins.",
            "Agent JWTs use RS256 with kid. Private keys remain backend-only. Public keys are versioned.",
            "JWT lifetime defaults to 60 minutes. Revocation stores jti and agent_id state.",
            "Key rotation keeps old public keys valid until all issued tokens expire, then disables the old kid.",
            "Emergency rotation invalidates active jti rows and marks affected agents as suspended until respawned.",
        ]),
        h(1, "5. Ledger Integrity Design"),
        numbered([
            "The first row uses a fixed genesis prev_hash value stored in configuration and documented in migration history.",
            "Each row hash is SHA-256 over canonical JSON event_data, prev_hash, tenant_id, agent_id, event_type, verdict, severity, and created_at.",
            "Ledger append occurs in the same transaction as verdict state changes where possible.",
            "Concurrent writes use a PostgreSQL advisory lock or serializable transaction to ensure exactly one previous head.",
            "Application DB role can INSERT and SELECT audit_ledger but cannot UPDATE or DELETE.",
            "A database trigger rejects any UPDATE or DELETE attempt even if a broader role is accidentally used.",
            "Verification scans rows in id order and returns valid, entries_checked, first broken id, expected hash, and actual hash.",
        ]),
        h(1, "6. Trust Score Rules"),
        table([
            ["Event", "Trust Delta", "Notes"],
            ["Clean allowed message", "+0.005", "Cap positive movement to once per minute per agent."],
            ["Low-confidence suspicious content flagged", "-0.03", "Verdict FLAGGED unless policy says block."],
            ["High-confidence injection blocked", "-0.15", "Verdict BLOCKED; create threat event."],
            ["Unauthorized tool call blocked", "-0.20", "High severity if tool is sensitive."],
            ["Invalid or spoofed identity", "-0.30", "Critical; may suspend agent depending policy."],
            ["Manual resolve false positive", "+0.05", "Audited admin action; never deletes event."],
        ]),
        p("Trust scores are bounded between 0.0 and 1.0. The first production version uses trust for visibility and risk escalation, not automatic permission expansion. Adaptive permission escalation is deferred until policy review exists."),
        h(1, "7. Security Acceptance Gates"),
        table([
            ["Gate", "Pass Condition"],
            ["Auth", "Missing API key, invalid API key, expired JWT, mismatched agent_id, and revoked jti all fail with expected error codes."],
            ["Policy", "Agent cannot call a tool/action absent from manifest."],
            ["Ledger", "Tampered row causes /ledger/verify to return valid=false at the modified row."],
            ["Latency", "Synchronous deterministic path meets target without invoking LLM."],
            ["Secrets", "No raw API keys, private JWT keys, or prompt payload secrets appear in frontend bundle or logs."],
        ]),
        h(1, "8. Incident Response Runbooks"),
        table([
            ["Incident", "Immediate Action", "Follow-Up"],
            ["API key compromise", "Disable key status, identify last_used_at and request logs, create replacement key.", "Review tenant events, rotate any linked demo credentials, document blast radius."],
            ["JWT private key compromise", "Disable current kid, revoke active jti rows, suspend affected agents, rotate private/public key pair.", "Audit token issuance logs and require agent respawn."],
            ["Ledger verification failure", "Stop protected write path, snapshot database, identify first broken row, compare backup or export.", "Treat as security incident until explained by migration/test data."],
            ["Critical threat spike", "Rate limit affected tenant, surface dashboard alert, collect sample evidence.", "Tune detector patterns and update curated attack set."],
            ["LLM enrichment outage", "Keep deterministic protection online, mark enrichment degraded in health status.", "Retry queued jobs after provider recovery."],
        ]),
        h(1, "9. Residual Risk Register"),
        table([
            ["Risk", "Why It Remains", "Owner Action"],
            ["No execution sandbox", "Original scope excludes isolation of arbitrary tool execution.", "Document integration boundary and require host application sandboxing for dangerous tools."],
            ["Pattern evasion", "Attackers can invent new phrasing that avoids curated regexes.", "Maintain labeled dataset and use async enrichment to discover new signatures."],
            ["Operator misuse", "Admins may create overly broad manifests.", "Add risky-permission warnings and require review for write/destructive tools."],
            ["Database superuser tampering", "Application-level append-only controls cannot stop a fully compromised DB superuser.", "Use backups, external hash anchoring, and restricted production DB access."],
        ]),
    ],
)


RUNBOOK = DocSpec(
    filename="AgentShield_Implementation_Runbook.docx",
    title="AgentShield Implementation Runbook",
    subtitle="Step-by-step production build order for backend, frontend, SDK, security engine, deployment, and observability.",
    blocks=[
        h(1, "1. Repository Layout"),
        table([
            ["Path", "Purpose"],
            ["backend/app", "FastAPI application, routers, services, middleware."],
            ["backend/app/security", "API key auth, JWT identity, permission engine, injection detector."],
            ["backend/app/ledger", "Ledger append, canonical hash, verification, DB triggers."],
            ["backend/app/events", "Outbox, WebSocket manager, event serializers."],
            ["backend/migrations", "Alembic migrations for PostgreSQL."],
            ["frontend/app", "Next.js App Router pages."],
            ["frontend/components", "Dashboard, ledger, attack sim, shared UI, Three.js scene."],
            ["sdk/python/agentshield", "Python SDK package."],
            ["tests", "Backend, SDK, contract, and e2e tests."],
        ]),
        h(1, "2. Phase 0 - Foundation"),
        numbered([
            "Create backend FastAPI project with health endpoint, settings loader, structured logging, and test harness.",
            "Create PostgreSQL migrations for tenants, api_keys, agents, agent_tokens, permission_manifests, audit_ledger, threat_events, trust_history, and event_outbox.",
            "Create frontend Next.js app with Tailwind, base layout, dashboard route shell, and API client placeholder.",
            "Create SDK package skeleton with typed client, exceptions, and local example script.",
            "Add CI commands for backend tests, frontend lint/build, SDK tests, and contract schema checks.",
        ]),
        table([
            ["Acceptance", "Verification Command"],
            ["Backend boots and /health returns ok.", "uvicorn app.main:app --reload; curl /health"],
            ["Migrations apply from empty database.", "alembic upgrade head"],
            ["Frontend builds empty shells.", "npm run build"],
            ["SDK imports.", "python -c 'import agentshield'"],
        ]),
        h(1, "3. Phase 1 - Security Core"),
        numbered([
            "Implement API key generation, hashing, lookup, scopes, and auth middleware.",
            "Implement RS256 key loading, JWT issuance, validation, kid support, jti registry, expiry, and revocation.",
            "Implement permission manifests with deny-by-default semantics and action-level checks.",
            "Implement deterministic injection detector using curated regexes, normalized text, confidence scoring, and evidence spans.",
            "Implement /v1/agents, /v1/shield/analyze, and /v1/shield/tool-call using documented schemas.",
        ]),
        table([
            ["Acceptance", "Verification Command"],
            ["Invalid credentials are rejected.", "pytest tests/security/test_auth.py"],
            ["Unauthorized tool calls are blocked.", "pytest tests/security/test_permissions.py"],
            ["Known injection payloads are blocked or flagged.", "pytest tests/security/test_injection_detector.py"],
        ]),
        h(1, "4. Phase 2 - Ledger, Trust, And Events"),
        numbered([
            "Implement canonical JSON hash function and genesis hash constant.",
            "Implement ledger append service with transaction lock and immutable database trigger migration.",
            "Implement trust score delta rules and trust_history writes.",
            "Implement event_outbox and WebSocket broadcast for security events, trust updates, and ledger verification.",
            "Implement /v1/ledger, /v1/ledger/verify, /v1/threats, and trust history endpoints.",
        ]),
        table([
            ["Acceptance", "Verification Command"],
            ["Ledger writes are hash chained.", "pytest tests/ledger/test_append.py"],
            ["Concurrent writes preserve chain.", "pytest tests/ledger/test_concurrency.py"],
            ["Tamper simulation breaks verification.", "pytest tests/ledger/test_verify.py"],
            ["WebSocket receives committed events.", "pytest tests/events/test_websocket.py"],
        ]),
        h(1, "5. Phase 3 - SDK And Demo Agent"),
        numbered([
            "Implement Python SDK client with retries, timeout, typed responses, and SecurityBlocked exception.",
            "Create demo research agent flow that analyzes user messages before processing and checks tool calls before execution.",
            "Create 10 curated attack fixtures covering instruction override, role hijack, data exfiltration, tool misuse, and spoofing.",
            "Implement /v1/attack-sim/run using the real analyze and tool-call paths; do not mock verdicts.",
        ]),
        h(1, "6. Phase 4 - Frontend Dashboard"),
        numbered([
            "Build dashboard route with live metrics, WebSocket connection state, event feed, and agent trust panel.",
            "Build ledger route with paginated table, full hash expansion, verify-chain button, and client-side tamper demo clearly labeled as local simulation.",
            "Build attack simulator route with attack selector, custom payload editor, run button, verdict panel, evidence, and latency.",
            "Build agent registry route with permissions, status, trust history, revoke action, and empty/loading/error states.",
            "Add Three.js network only after REST/WebSocket payloads are stable; use list fallback for mobile.",
        ]),
        table([
            ["Acceptance", "Verification Command"],
            ["Dashboard connects to backend and updates from WebSocket.", "Playwright dashboard smoke test."],
            ["Ledger page verifies real chain.", "Playwright ledger smoke test."],
            ["Attack sim blocks curated payload.", "Playwright attack simulation smoke test."],
            ["Mobile fallback is usable.", "Playwright mobile viewport smoke test."],
        ]),
        h(1, "7. Phase 5 - Deployment And Operations"),
        numbered([
            "Provision Railway backend and PostgreSQL with separate app and migration credentials.",
            "Provision Vercel frontend with NEXT_PUBLIC_API_URL and NEXT_PUBLIC_WS_URL.",
            "Configure secrets: DATABASE_URL, JWT_PRIVATE_KEY, JWT_PUBLIC_KEYS_JSON, API_KEY_PEPPER, OPENAI_API_KEY for optional enrichment, CORS origins.",
            "Run migrations, seed tenant and demo API key, spawn demo agents, and record non-secret identifiers.",
            "Run production smoke: health, agent spawn, analyze allowed message, analyze blocked injection, tool-call denial, ledger verify, dashboard route, WebSocket event.",
        ]),
        h(1, "8. Observability Requirements"),
        bullets([
            "Log request_id, tenant_id, agent_id, endpoint, verdict, latency_ms, and ledger_id; never log raw API keys, private keys, or full sensitive messages by default.",
            "Expose metrics for request count, blocked count, latency, ledger append failures, WebSocket clients, queue depth, and enrichment failures.",
            "Alert on ledger append failure, ledger verification failure, auth spike, critical threat spike, and queue backlog.",
            "Provide an incident runbook for key compromise, ledger integrity failure, and backend outage.",
        ]),
        h(1, "9. Definition Of Done For Each Phase"),
        table([
            ["Phase", "Done Means"],
            ["0", "A new developer or AI agent can clone the repo, configure env from example files, run migrations, run tests, and boot both apps locally."],
            ["1", "The server rejects invalid auth and blocks deterministic policy violations without any frontend or LLM dependency."],
            ["2", "Every protected decision has a ledger entry, trust changes are explainable, and event streaming works from committed outbox rows."],
            ["3", "The SDK examples protect a real demo flow and failures are ergonomic enough for application developers to handle."],
            ["4", "The dashboard is useful with real backend data, with no hidden mock state outside an explicitly labeled demo mode."],
            ["5", "A production operator can deploy, verify, monitor, rotate keys, recover from common incidents, and roll back safely."],
        ]),
        h(1, "10. AI Agent Build Rules"),
        bullets([
            "Do not implement dashboard animations before backend contracts and fixture payloads are stable.",
            "Do not put LLM calls in the default synchronous enforcement path.",
            "Do not store raw API keys, JWT private keys, or full sensitive payloads in logs or frontend state.",
            "Do not allow a successful protected action unless the ledger write succeeded.",
            "Do not treat the Three.js graph as required for mobile acceptance; the mobile list fallback is a first-class interface.",
            "When uncertain, choose the stricter security behavior and document it in the handoff.",
        ]),
    ],
)


TESTS = DocSpec(
    filename="AgentShield_Test_and_Acceptance_Plan.docx",
    title="AgentShield Test And Acceptance Plan",
    subtitle="Unit, integration, security, latency, ledger-integrity, browser, and production smoke tests.",
    blocks=[
        h(1, "1. Test Strategy"),
        p("AgentShield must be tested as a security control, not only as a web app. The core acceptance standard is that protected actions are either allowed or blocked for explicit, auditable reasons, and every decision can be verified later."),
        h(1, "2. Unit Test Matrix"),
        table([
            ["Area", "Scenarios"],
            ["API key auth", "missing, malformed, inactive, wrong scope, expired, rate limited, valid."],
            ["JWT identity", "valid, expired, wrong issuer, wrong audience, wrong agent_id, revoked jti, rotated kid."],
            ["Permission engine", "allowed action, missing tool, missing action, wildcard disallowed, denied default."],
            ["Injection detector", "known injection, benign request, mixed benign/malicious, casing/spacing obfuscation, data exfil phrase."],
            ["Trust scoring", "delta bounds, repeated clean cap, blocked threat penalty, manual resolve recovery."],
            ["Ledger hashing", "canonical ordering, genesis row, changed event_data, changed prev_hash, changed timestamp."],
        ]),
        h(1, "3. Integration Test Matrix"),
        table([
            ["Flow", "Expected Result"],
            ["Spawn agent then analyze benign message.", "ALLOWED, trust unchanged or small positive, ledger row written."],
            ["Analyze high-confidence injection.", "BLOCKED, CRITICAL or HIGH, threat event written, trust reduced, ledger row written."],
            ["Check unauthorized tool call.", "BLOCKED with POLICY_TOOL_DENIED, trust reduced, ledger row written."],
            ["Use expired token.", "401 AUTH_AGENT_TOKEN_INVALID, auth ledger event if configured."],
            ["Verify ledger after normal writes.", "valid=true and entries count matches."],
            ["Tamper a copied test database row.", "valid=false and broken_at equals changed row."],
            ["WebSocket client connected during attack sim.", "Receives security.event.created and attack_sim.completed."],
        ]),
        h(1, "4. Security Test Cases"),
        table([
            ["Test", "Pass Condition"],
            ["No raw secrets in logs", "Search logs for API key prefix, JWT private key markers, and OPENAI_API_KEY returns no leaks."],
            ["Frontend bundle secret scan", "No private env vars or raw keys appear in built JS."],
            ["CORS enforcement", "Unauthorized origins fail browser preflight for protected requests."],
            ["Rate limit", "Burst above configured threshold returns RATE_LIMITED."],
            ["Body size limit", "Oversized payload rejected before classifier work."],
            ["Mutation protection", "Application DB role cannot update or delete audit_ledger."],
        ]),
        h(1, "5. Latency And Load Acceptance"),
        table([
            ["Endpoint", "Target", "Conditions"],
            ["POST /v1/shield/analyze deterministic path", "P95 < 200 ms", "No LLM enrichment in request path; includes ledger write."],
            ["POST /v1/shield/tool-call", "P95 < 150 ms", "Permission check, trust update, ledger write."],
            ["GET /v1/ledger", "P95 < 300 ms", "Paginated first page at 10k rows."],
            ["GET /v1/ledger/verify", "Completes for 10k rows under operational threshold", "May be async for larger ledgers."],
            ["WebSocket event delivery", "Visible in dashboard under 1 second", "After backend commit."],
        ]),
        h(1, "6. Browser Acceptance"),
        table([
            ["Route", "Desktop Acceptance", "Mobile Acceptance"],
            ["/dashboard", "Shows live counters, event feed, 3D graph or initialized scene, agent panel.", "Shows list fallback, counters, event feed, no horizontal overflow."],
            ["/ledger", "Rows load, hash expands, verify action returns result.", "Rows are readable in compact layout."],
            ["/attack-sim", "Curated attack runs and produces blocked verdict with evidence.", "Payload editor and verdict remain usable."],
            ["/agents", "Agent list, trust history, permissions, revoke action states.", "Cards or compact rows fit viewport."],
        ]),
        h(1, "7. Production Smoke Checklist"),
        numbered([
            "GET /health returns ok from deployed backend.",
            "Frontend loads deployed dashboard without console errors.",
            "Spawn demo agent using production API key.",
            "Analyze benign message and confirm ALLOWED verdict plus ledger id.",
            "Analyze curated injection and confirm BLOCKED verdict plus threat event.",
            "Run unauthorized tool-call check and confirm BLOCKED verdict.",
            "GET /v1/ledger/verify returns valid=true.",
            "Open dashboard and confirm the new events appear through WebSocket or fallback polling.",
            "Run attack simulator from browser and confirm result matches backend ledger.",
        ]),
        h(1, "8. Release Gates"),
        table([
            ["Gate", "Required Evidence"],
            ["Contracts frozen", "API and event contract doc matches OpenAPI/schema files."],
            ["Core tests pass", "Backend unit/integration/security tests green."],
            ["Frontend checks pass", "Lint/build and browser smoke tests green."],
            ["Ledger proven", "Append, tamper, and concurrency tests green."],
            ["Production smoke pass", "Recorded endpoint responses and dashboard evidence."],
            ["Rollback known", "Documented migration rollback or recovery path for failed release."],
        ]),
        h(1, "9. Curated Attack Dataset Requirements"),
        table([
            ["Category", "Minimum Cases", "Expected Behavior"],
            ["Instruction override", "10", "High-confidence examples BLOCKED; ambiguous examples FLAGGED."],
            ["Role hijack", "8", "Attempts to redefine the agent role are BLOCKED or FLAGGED with evidence span."],
            ["Data exfiltration", "8", "Requests to reveal secrets, prompts, keys, or hidden context are BLOCKED."],
            ["Tool misuse", "8", "Requests for destructive or unauthorized actions are BLOCKED by permission policy."],
            ["Benign controls", "20", "Normal task requests remain ALLOWED to track false positives."],
        ]),
        h(1, "10. Acceptance Report Format"),
        bullets([
            "Record git commit, deployment URL, backend version, frontend version, database migration head, and test timestamp.",
            "Include command outputs or links for backend tests, frontend build, browser smoke, and production smoke.",
            "List skipped checks with a concrete reason and risk, not just 'not run'.",
            "Attach the final ledger verification result and at least one blocked attack verdict with ledger_id.",
            "Record any production defects found during smoke testing as follow-up issues before launch approval.",
        ]),
    ],
)


DOCS = [MASTER, CONTRACTS, SECURITY, RUNBOOK, TESTS]


def main() -> None:
    DOCX_DIR.mkdir(parents=True, exist_ok=True)
    MD_DIR.mkdir(parents=True, exist_ok=True)
    for spec in DOCS:
        build_docx(spec)
        build_md(spec)
    manifest = OUT_DIR / "README.md"
    manifest.write_text(
        "# AgentShield Production Documentation Pack\n\n"
        "This package supersedes the original AgentShield PRD, TRD, and design document as the production planning source of truth.\n\n"
        "## DOCX Deliverables\n"
        + "\n".join(f"- docx/{spec.filename}" for spec in DOCS)
        + "\n\n## Markdown Sources\n"
        + "\n".join(f"- sources/{spec.filename.replace('.docx', '.md')}" for spec in DOCS)
        + "\n\nOriginal files were not modified.\n",
        encoding="utf-8",
    )
    print(f"Generated {len(DOCS)} DOCX files and Markdown sources in {OUT_DIR}")


if __name__ == "__main__":
    main()
